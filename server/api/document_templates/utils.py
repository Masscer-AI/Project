"""
DOCX placeholder extraction for Jinja-style ``{{ variable_name }}`` tags.

Used at upload time and when replacing the template file to refresh
``metadata["placeholders"]`` and merge ``metadata["variables"]``.
"""

from __future__ import annotations

import os
import re
import tempfile
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# Match simple Jinja placeholders (same spirit as docussistant server regex).
PLACEHOLDER_RE = re.compile(r"\{\{\s*([^\{\}\s]+)\s*\}\}")


def _findall_ordered_unique(text: str) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for name in PLACEHOLDER_RE.findall(text or ""):
        if name not in seen:
            seen.add(name)
            out.append(name)
    return out


def _iter_block_items(parent):
    """Yield paragraphs and tables in document body order."""
    for child in parent.element.body.iterchildren():
        if child.tag.endswith("tbl"):
            yield Table(child, parent)
        elif child.tag.endswith("p"):
            yield Paragraph(child, parent)


def extract_placeholders_from_docx_path(path: str) -> list[str]:
    """
    Extract placeholder names from a .docx file (body, tables, headers, footers).

    Args:
        path: Local filesystem path readable by python-docx.

    Returns:
        Ordered unique placeholder names.
    """
    doc = Document(path)
    names: list[str] = []
    seen: set[str] = set()

    def add_many(text: str) -> None:
        nonlocal names, seen
        for n in _findall_ordered_unique(text):
            if n not in seen:
                seen.add(n)
                names.append(n)

    # Body
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            add_many(block.text)
        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        add_many(p.text)

    # Headers / footers (all sections)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                add_many(p.text)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            add_many(p.text)

    return names


def extract_placeholders_from_storage_file(file_field) -> list[str]:
    """
    Read bytes from a Django ``FieldFile`` and extract placeholders (S3-safe).
    """
    with file_field.open("rb") as f:
        raw = f.read()
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    try:
        tmp.write(raw)
        tmp.close()
        return extract_placeholders_from_docx_path(tmp.name)
    finally:
        try:
            os.unlink(tmp.name)
        except OSError:
            pass


def merge_variables_metadata(
    previous_variables: dict | None,
    placeholders: list[str],
) -> dict[str, dict]:
    """
    Build variables map for ``DocumentTemplate.metadata["variables"]``.

    Preserves description/required/example for placeholders that still exist;
    creates draft entries for new placeholders; drops keys not in ``placeholders``.
    """
    prev = previous_variables or {}
    out: dict[str, dict] = {}
    for ph in placeholders:
        old = prev.get(ph) if isinstance(prev.get(ph), dict) else {}
        out[ph] = {
            "description": str(old.get("description", "") or ""),
            "required": bool(old.get("required", True)),
            "example": str(old.get("example", "") or ""),
        }
    return out


def build_template_metadata(
    placeholders: list[str],
    previous_metadata: dict | None = None,
) -> dict:
    """Full metadata dict for a template after extraction."""
    prev = (previous_metadata or {}) if isinstance(previous_metadata, dict) else {}
    prev_vars = prev.get("variables")
    if not isinstance(prev_vars, dict):
        prev_vars = {}
    return {
        "placeholders": list(placeholders),
        "variables": merge_variables_metadata(prev_vars, placeholders),
    }
