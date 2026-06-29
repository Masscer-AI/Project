import fitz
import chardet
from docx import Document as DocxDocument
from io import BytesIO
from .models import Chunk, Document, Collection
from api.utils.color_printer import printer
from api.utils.openai_functions import (
    create_structured_completion,
    create_completion_openai,
)
from pydantic import BaseModel, Field


def detect_file_encoding(file):
    raw_data = file.read(10000)
    result = chardet.detect(raw_data)
    file.seek(0)
    return result["encoding"]


def _read_file_head(file, n: int = 8) -> bytes:
    pos = file.tell()
    head = file.read(n)
    file.seek(pos)
    return head


def infer_upload_format(
    file,
    *,
    content_type: str | None = None,
    fallback_name: str | None = None,
) -> str:
    """
    Resolve the upload format from filename, MIME type, and magic bytes.

    Production uploads sometimes arrive without a useful extension; MIME and
    file signatures are used as fallbacks. ``fallback_name`` (e.g. form field
    ``name``) is used when the multipart filename is missing or has no suffix.
    """
    name = getattr(file, "name", "") or ""
    if fallback_name and ("." not in name or not name.strip()):
        name = fallback_name
    extension = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    ctype = (
        (content_type or getattr(file, "content_type", "") or "")
        .lower()
        .split(";")[0]
        .strip()
    )

    if extension in ("pdf", "docx", "xlsx", "xlsm", "xls", "txt", "html", "csv"):
        return extension

    if ctype in ("application/pdf",) or ctype.endswith("/pdf"):
        return "pdf"
    if "wordprocessingml" in ctype:
        return "docx"
    if "spreadsheetml" in ctype or ctype in (
        "application/vnd.ms-excel",
        "application/msexcel",
    ):
        return "xlsx"
    if ctype.startswith("text/"):
        return "html" if "html" in ctype else "txt"

    head = _read_file_head(file, 5)
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"PK\x03\x04"):
        lower_name = name.lower()
        if any(token in lower_name for token in ("xls", "sheet", "excel")):
            return "xlsx"
        if "doc" in lower_name:
            return "docx"
        if "spreadsheet" in ctype or "excel" in ctype:
            return "xlsx"
        if "wordprocessing" in ctype or "msword" in ctype:
            return "docx"
        return "office_zip"

    return extension or "txt"


def _read_xlsx_content(raw: bytes, file_name: str) -> tuple[str, str]:
    try:
        from api.utils.spreadsheet_tools import extract_xlsx_text_from_bytes
    except ModuleNotFoundError as exc:
        raise ValueError(
            "Excel support is not installed on this server (missing openpyxl). "
            "Rebuild and redeploy the Django image."
        ) from exc

    try:
        text = extract_xlsx_text_from_bytes(raw)
    except Exception as exc:
        raise ValueError(f"Could not read Excel file: {exc}") from exc
    return text, file_name


def _read_office_zip_content(raw: bytes, file_name: str) -> tuple[str, str]:
    try:
        from api.utils.spreadsheet_tools import extract_xlsx_text_from_bytes
    except ModuleNotFoundError as exc:
        raise ValueError(
            "Excel support is not installed on this server (missing openpyxl). "
            "Rebuild and redeploy the Django image."
        ) from exc

    try:
        return extract_xlsx_text_from_bytes(raw), file_name
    except ValueError:
        raise
    except Exception:
        doc = DocxDocument(BytesIO(raw))
        text = "\n".join([para.text for para in doc.paragraphs])
        if not text.strip():
            raise ValueError(
                "Could not read Office file. Rename it to .xlsx or .docx and try again."
            )
        return text, file_name


def read_file_content(
    file,
    content_type: str | None = None,
    fallback_name: str | None = None,
):
    file_extension = infer_upload_format(
        file,
        content_type=content_type,
        fallback_name=fallback_name,
    )
    file_name = (fallback_name or "").strip() or file.name

    if file_extension == "pdf":
        raw = file.read()
        file.seek(0)
        doc = fitz.open(stream=raw, filetype="pdf")
        text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text()
        return text, file_name
    elif file_extension == "docx":
        raw = file.read()
        file.seek(0)
        doc = DocxDocument(BytesIO(raw))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text, file_name
    elif file_extension in ("xlsx", "xlsm"):
        raw = file.read()
        file.seek(0)
        return _read_xlsx_content(raw, file_name)
    elif file_extension == "office_zip":
        raw = file.read()
        file.seek(0)
        try:
            return _read_office_zip_content(raw, file_name)
        except Exception as exc:
            raise ValueError(
                "Could not read Office file. Rename it to .xlsx or .docx and try again."
            ) from exc
    elif file_extension == "xls":
        raise ValueError(
            "Legacy .xls files are not supported. Save the file as .xlsx and try again."
        )
    else:
        head = _read_file_head(file, 4)
        if head.startswith(b"PK\x03\x04"):
            raw = file.read()
            file.seek(0)
            return _read_office_zip_content(raw, file_name)

        file_encoding = detect_file_encoding(file)
        try:
            text = file.read().decode(file_encoding)
        except UnicodeDecodeError as exc:
            raise ValueError(
                "Could not decode file as text. If this is an Excel file, "
                "ensure it is saved as .xlsx and uploaded with that extension."
            ) from exc
        file.seek(0)
        return text, file_name


class ChunkBriefModel(BaseModel):
    tags: str = Field(
        ...,
        description="A comma-separated list of tags associated with the chunk. Up to 3 words",
    )
    brief: str = Field(
        ...,
        description="A brief summary of max 25 words that describes the chunk content",
    )


def generate_chunk_brief(chunk_id):
    _system = """
You are a Machine Learning an AI expert. You are classifying chunks for a RAG pipeline. The goals is tu summarice and classify the content of an abstract piece of text. You goal is to return kwywords that can be potencially related to the chunk content and to summarice its content in up to 80 words, try to explain briefly what is the content of the chunk.


Everything must be in the same language as the chunk.
"""

    c = Chunk.objects.get(pk=chunk_id)
    printer.red(f"Generating chunk brief for {c.pk}")
    printer.blue(c.content)
    # TODO: register consumption here
    chunkito = create_structured_completion(
        model="gpt-4o-mini",
        system_prompt=_system,
        user_prompt=c.content,
        response_format=ChunkBriefModel,
    )
    printer.green(chunkito)
    c.brief = chunkito.brief
    c.tags = chunkito.tags
    c.save()
    return True


def generate_document_brief(document_id: int):
    number_of_characters = 50000
    _system = f"""
You are an AI and Machine Learning Specialist.
You task is to summarice the content of a document in up to 100 words, explaining what is the document and an overview of its content.
You will receive the first {number_of_characters} characters in the document.

The brief must be in the same language as the document.
"""
    d = Document.objects.get(pk=document_id)
    first_20000_chars = d.text[:number_of_characters]
    # TODO: register consumption here
    brief = create_completion_openai(
        system_prompt=_system, user_message=first_20000_chars
    )
    d.brief = brief
    d.save()


class SelectedChunks(BaseModel):
    queries: list[str] = Field(
        ...,
        description="A list of 4 different queries that can lead to the best results when using a vector store to retrieve information about the collection where the chunks are stored. The goals is to provide another AI with relevant context from the vector storage",
    )
    tags: str = Field(
        ...,
        description="A command separated list of related tags",
    )
    search_string: str = Field(
        ...,
        description="A one up to three words string that can be in the texts stored in the vector storage system",
    )


def querify_context(context: str) -> SelectedChunks:
    # TODO: Instead, get the collection summary to understand the context better
    # chunks = get_chunks_for_collection(collection)
    # printer.yellow(len(chunks), "Number of chunks for the collection")
    # chunks_str = " ".join([json.dumps({"brief": c.content[100:500]}) for c in chunks[:300]])

    _system = f"""
You are a AI and Machine Learning specialist.

The following context takes part of a conversation between the user and an AI assistant. Your task is to create queries that can lead to the best results when querying the vector storage as per user requirement.

CONVERSATION CONTEXT:
---
{context}
---

"""
    # TODO: register consumption here
    queries = create_structured_completion(
        model="gpt-4o-mini",
        system_prompt=_system,
        user_prompt="Return the most accurate and relevant queries to retrieve information from the vector storage, making mention of important keywords for each one.",
        response_format=SelectedChunks,
    )
    return queries


def get_chunks_for_collection(collection):
    # Get all documents associated with the collection
    documents = Document.objects.filter(collection=collection)

    # Get all chunks associated with those documents
    chunks = Chunk.objects.filter(document__in=documents)

    return chunks


def extract_rag_results(rag_results, context):
    documents_context = ""
    complete_context = context
    counter = 0
    added_sources = []
    sources = []
    if rag_results is not None:
        metadatas = rag_results["results"]["metadatas"]
        for meta in metadatas:
            for ref in meta:
                if len(ref) > 0:
                    if ref in added_sources:
                        continue

                    sources.append(ref)
                    added_sources.append(ref)

                    documents_context += f"vector_href='#{ref.get('model_name', 'chunk')}-{ref.get('model_id', 132132)}'\nVECTOR CONTENT:\n--- {ref.get('content', '')}---"
                    counter += 1

        if len(documents_context) > 0:
            complete_context += f"\n\nThe following is information about a embeddings vector storage querying the user message: ---start_vector_context\n\n{documents_context}\n\n---end_vector_context---\nIf you use information from the vector storage, please cite the resourcess in anchor tags using the provided href, for example: <a href='#chunk-CHUNK_ID' target='__blank'>SOME_RELATED_CONECTOR</a> where  SOME_RELATED_CONECTOR is a three-four words text related to the chunk content that the user will be able to review. You can add the sources in any place of your response. Add as many as needed. You must cite the source using the href, the SOME_RELATED_CONECTOR is generated by you."

    return complete_context, sources
